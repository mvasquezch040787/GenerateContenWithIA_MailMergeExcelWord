import customtkinter.windows
import google.generativeai as genai
from Data import Texts
import customtkinter 
import tkinter as tk
import openpyxl
from tqdm import tqdm
import time
from tkinter import ttk
from docx import Document
import pandas as pd

workBookData=""
sheetTopics=""
sheetData = ""
customtkinter.set_appearance_mode("light")
app = customtkinter.CTk()

app.title("Generate")

lblVacio=customtkinter.CTkLabel(master=app, text="          ")
lblVacio.grid(row=0,column=0)

lblCarrera=customtkinter.CTkButton(master=app, text="Carrera",fg_color="gray")
lblCarrera.grid(row=0,column=1)

TxtCarrera=customtkinter.CTkEntry(master=app, width=1250)
TxtCarrera.grid(row=0,column=2)


lblUnidad=customtkinter.CTkButton(master=app, text="Unidad Didactica ",fg_color="gray")
lblUnidad.grid(row=1,column=1)

TxtUnidad=customtkinter.CTkEntry(master=app, width=1250,placeholder_text="Ingresar la Unidad Didactica")
TxtUnidad.grid(row=1,column=2)

TxtMessage=customtkinter.CTkEntry(master=app, width=1250)
TxtMessage.grid(row=5,column=2)

TxtCarrera.insert(0,"ARQUITECTURA DE PLATAFORMAS Y SERVICIOS DE TECNOLOGÍAS DE INFORMACIÓN")


def seleccionarTemasExcel():
    fileExcelTopics =tk.filedialog.askopenfilename(filetypes=[('Archivos Excel', '*.xlsx')])
    workbook = openpyxl.load_workbook(fileExcelTopics)
    global sheetTopics 
    sheetTopics=workbook.active 

def seleccionarArchivoData():
    global fileExcelData
    fileExcelData = tk.filedialog.askopenfilename(filetypes=[('Archivos Excel', '*.xlsx')])
    global workBookData
    workBookData=openpyxl.load_workbook(fileExcelData)
    global sheetData
    sheetData=workBookData.active

def seleccionarArchivoTemplateWord():
    global fileTemplateWord
    fileTemplateWord = tk.filedialog.askopenfilename(filetypes=[('Archivos Word', '*.docx')])

def generateMailMerge():
    docxTemplate=Document(fileTemplateWord)
    ''' dataExcel=pd.read_excel(fileExcelData)'''
    
    marcadores=[]
    for paragraph in docxTemplate.paragraphs:
        for run in paragraph.runs:
            if run.text.startswith('<<') and run.text.endswith('>>'):
                marcador = run.text[2:-2]  # Elimina los delimitadores << y >>
                marcadores.append(marcador)
    print(docxTemplate)
    print(marcadores)

'''
    for index, fila in dataExcel.iterrows():
        for parrafo in docxTemplate.paragraphs:
            # Clonar el párrafo original (con su formato)
            nuevo_parrafo = documentoNew.add_paragraph()
            for run in parrafo.runs:
                texto = run.text
                texto = texto.replace('{AA}', str(fila['NAA']))
                texto = texto.replace('{IL}', str(fila['ILS']))
                texto = texto.replace('{SS}', str(fila['NS']))
                texto = texto.replace('{TS}', str(fila['TEMA']))
                texto = texto.replace('{PS}', str(fila['PROPOSITO']))
                texto = texto.replace('{SA}', str(fila['TEMA']))
                texto = texto.replace('{RA}', str(fila['RD']))
                texto = texto.replace('{IN}', str(fila['INICIO']))
                texto = texto.replace('{DE}', str(fila['DESARROLLO']))
                texto = texto.replace('{CI}', str(fila['CIERRE']))
                
                # Agregar texto y formato al nuevo párrafo
                nuevo_run = nuevo_parrafo.add_run(texto)
                nuevo_run.bold = run.bold
                nuevo_run.italic = run.italic
                nuevo_run.underline = run.underline
                nuevo_run.font.size = run.font.size
                nuevo_run.font.name = run.font.name
    
    # Guardar el documento modificado
    documentoNew.save('PC_COMBINADA.DOCX')
'''

def Generate():
    genai.configure(api_key=Texts.API_KEY)
    model=genai.GenerativeModel(model_name="gemini-pro")    
    row_count=2
    for row in sheetTopics.iter_rows(min_row=1,max_row=16,max_col=1):
        for cell in row:
            try:
                if cell.value is None:
                    break
                else:
                    time.sleep(1.5)
                    IL=str(model.generate_content(Texts.ACCURACYIL+cell.value+ " De "+TxtUnidad.get()+ "de la carrera de "+TxtCarrera.get()).text)
                    time.sleep(1.5)
                    sheetData[f'A{row_count}']=row_count-1
                    sheetData[f'B{row_count}']=IL
                    sheetData[f'C{row_count}']=row_count-1
                    sheetData[f'D{row_count}']=cell.value
                    sheetData[f'E{row_count}']=(model.generate_content(Texts.ACCURACY_PROPOSITO+cell.value+ " De "+IL).text)
                    time.sleep(1.5)
                    sheetData[f'F{row_count}']=(model.generate_content(Texts.ACCURACYIL_RECURSOS+cell.value+ " De "+IL).text)
                    time.sleep(1.5)
                    sheetData[f'G{row_count}']=(model.generate_content(Texts.ACCURACY_INICIO+cell.value+ " De "+IL).text)
                    time.sleep(1.5)
                    sheetData[f'H{row_count}']=(model.generate_content(Texts.ACCURACY_DESARROLLO+cell.value+ " De "+IL).text)
                    time.sleep(1.5)
                    sheetData[f'I{row_count}']=(model.generate_content(Texts.ACCURACY_CIERRE+cell.value+ " De "+IL).text)
                    time.sleep(1.5)
                    workBookData.save(fileExcelData)
                    time.sleep(1.5)
                    row_count+=1
            except ValueError as e:
                print(e)
        if row_count>17:
            TxtMessage.insert(tk.END, f"FELICITACIONES GENERACIÒN SATISFACTORIA")
            break
  
btnTopics=customtkinter.CTkButton(master=app, text="Seleccionar Temas EXCEL", width=1250,command=seleccionarTemasExcel)
btnTopics.grid(row=2,column=2)

BtnChangeXlsx=customtkinter.CTkButton(master=app, text="Seleccionar Almacenamiento EXCEL",width=1250,command=seleccionarArchivoData)
BtnChangeXlsx.grid(row=3,column=2)

BtnGenerate=customtkinter.CTkButton(master=app, text="GENERAR CONTENIDO USANDO IA",width=1250,command=Generate)
BtnGenerate.grid(row=4,column=2)

BtnChangeDocx=customtkinter.CTkButton(master=app, text="Seleccionar Plantilla WORD",width=1250,command=seleccionarArchivoTemplateWord)
BtnChangeDocx.grid(row=6,column=2)

BtnGenerateDocx=customtkinter.CTkButton(master=app, text="GENERAR DOCUMENTO WORD",width=1250,command=generateMailMerge)
BtnGenerateDocx.grid(row=7,column=2)

app.state("zoomed")
app.mainloop()




